#!/usr/bin/env python3
"""build_overview_docx.py — render docs/APP_FUNCTIONS_AND_DESIGN.md to a
styled Word document carrying the artifact's manuscript design (cream
page, warm ink, maroon headings with gold pane numbers, gold small-cap
Part eyebrows, pecha-style dividers, the maroon-ruled commitments
block). Parses the markdown so the docx always tracks the source.

Usage: python3 tools/build_overview_docx.py
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(HERE, "docs", "APP_FUNCTIONS_AND_DESIGN.md")
OUT = os.path.join(HERE, "docs", "APP_FUNCTIONS_AND_DESIGN.docx")

PAPER = "FAF6EE"
INK = RGBColor(0x2B, 0x21, 0x18)
INKSOFT = RGBColor(0x5C, 0x4F, 0x40)
ACCENT = RGBColor(0x8C, 0x2F, 0x2B)
GOLD = RGBColor(0x9A, 0x7A, 0x33)
RULE = RGBColor(0xC9, 0xB9, 0x92)
SERIF = "Palatino"
SANS = "Helvetica Neue"


def make_doc():
    doc = Document()
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), PAPER)
    doc.element.insert(0, bg)
    doc.settings.element.append(OxmlElement("w:displayBackgroundShape"))
    st = doc.styles["Normal"]
    st.font.name = SERIF
    st.font.size = Pt(11)
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(9)
    st.paragraph_format.line_spacing = 1.32
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1.15)
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    return doc


def add_run(p, text, size=None, color=None, bold=False, italic=False,
            font=None, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = font or SERIF
    if font:
        r._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    if caps:
        c = OxmlElement("w:caps")
        c.set(qn("w:val"), "1")
        r._element.rPr.append(c)
    if spacing:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        r._element.rPr.append(sp)
    return r


def rich(p, text):
    """inline **bold** / *italic* markup."""
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*|\*(.+?)\*", text, re.S):
        if m.start() > pos:
            add_run(p, text[pos:m.start()])
        if m.group(1) is not None:
            add_run(p, m.group(1), bold=True)
        else:
            add_run(p, m.group(2), italic=True)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:])


def para(doc, align=None, before=None, after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def pecha(doc):
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 18, 18)
    add_run(p, "─" * 22 + "  ", size=9, color=RULE)
    add_run(p, "❖", size=9, color=GOLD)
    add_run(p, "  " + "─" * 22, size=9, color=RULE)


def commitment(doc, text):
    p = para(doc)
    rich(p, text)
    pPr = p._element.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "8C2F2B")
    b.append(left)
    pPr.append(b)
    p.paragraph_format.left_indent = Inches(0.18)


def main():
    lines = open(MD, encoding="utf-8").read().split("\n")
    doc = make_doc()
    i = 0
    n = len(lines)

    def gather(j):
        """collect a blank-line-delimited block starting at j."""
        buf = []
        while j < n and lines[j].strip() != "" and not lines[j].startswith(
                ("#", "---", "|")):
            buf.append(lines[j])
            j += 1
        return " ".join(buf).strip(), j

    while i < n:
        ln = lines[i]
        s = ln.strip()
        if s == "" or s == "---":
            i += 1
            continue
        if s.startswith("# "):        # doc title
            title = s[2:]
            para(doc, WD_ALIGN_PARAGRAPH.CENTER, 40, 4)
            p = doc.paragraphs[-1]
            add_run(p, title, size=26, bold=True, color=ACCENT)
            i += 1
            continue
        if s.startswith("## A Complete Overview"):
            p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 0, 20)
            add_run(p, s[3:], size=13, italic=True, color=INKSOFT)
            pecha(doc)
            i += 1
            continue
        if s.startswith("# PART"):    # eyebrow + centered title follows
            eb = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 6, 4)
            m = re.match(r"# (PART [IVX]+)", s)
            add_run(eb, m.group(1) if m else s[2:], size=8, color=GOLD,
                    bold=True, font=SANS, caps=True, spacing=28)
            rest = s.split("—", 1)
            i += 1
            continue
        if s.startswith("## "):       # section head (centered) OR pane
            head = s[3:]
            if re.match(r"\d+\.\s", head) or head.startswith(
                    "The machinery"):
                # pane / machinery head: left, maroon, gold number
                p = para(doc, before=14, after=5)
                mnum = re.match(r"(\d+)\.\s+(.*)", head)
                if mnum:
                    add_run(p, mnum.group(1) + "  ", size=13, color=GOLD)
                    add_run(p, mnum.group(2), size=13, bold=True,
                            color=ACCENT)
                else:
                    add_run(p, head, size=13, bold=True, color=ACCENT)
                p.paragraph_format.keep_with_next = True
            else:
                p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 2, 12)
                add_run(p, head, size=16, bold=True, color=INK)
                p.paragraph_format.keep_with_next = True
            i += 1
            continue
        if s.startswith("#### "):     # h4
            p = para(doc, before=12, after=4)
            add_run(p, s[5:], size=12, bold=True, color=INK)
            p.paragraph_format.keep_with_next = True
            i += 1
            continue
        if re.match(r"^\d+\.\s", s):  # numbered list
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                m = re.match(r"^(\d+)\.\s+(.*)", lines[i].strip())
                j = i + 1
                cont = []
                while j < n and lines[j].startswith("   "):
                    cont.append(lines[j].strip())
                    j += 1
                text = " ".join([m.group(2)] + cont)
                p = para(doc, before=0, after=6)
                add_run(p, m.group(1) + ".  ", color=GOLD, bold=True)
                rich(p, text)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                i = j
            continue
        if s.startswith("- "):        # bullet list
            while i < n and lines[i].strip().startswith("- "):
                m = lines[i].strip()[2:]
                j = i + 1
                cont = []
                while j < n and lines[j].startswith("  ") and \
                        not lines[j].strip().startswith("- "):
                    cont.append(lines[j].strip())
                    j += 1
                p = para(doc, before=0, after=6)
                add_run(p, "•  ", color=RULE, bold=True)
                rich(p, " ".join([m] + cont))
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                i = j
            continue
        # a commitments numbered item rendered as a ruled block:
        # detect the three bold-led commitment paragraphs
        block, j = gather(i)
        if re.match(r"^\d+\.\s+\*\*", block) and (
                "binding layer" in block or "never guesses" in block or
                "never redone" in block):
            commitment(doc, re.sub(r"^\d+\.\s+", "", block))
        elif block.startswith("*Prepared with reverence"):
            p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 22)
            add_run(p, block.strip("*"), italic=True, color=INKSOFT)
        else:
            p = para(doc)
            rich(p, block)
        i = j

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
