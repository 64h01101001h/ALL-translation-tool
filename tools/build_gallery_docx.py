#!/usr/bin/env python3
"""build_gallery_docx.py — render the fifteen-pane screenshot gallery
as a styled Word document (docs/PANE_GALLERY.docx), same manuscript
design as the overview docx: cream page, maroon headings with gold
pane numbers, pecha divider, soft-ink captions. Images come from
docs/screenshots/ (regenerate them with the app's --screenshots mode).

Usage: python3 tools/build_gallery_docx.py
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(HERE, 'docs', 'screenshots')
OUT = os.path.join(HERE, 'docs', 'PANE_GALLERY.docx')

PAPER = 'FAF6EE'
INK = RGBColor(0x2B, 0x21, 0x18)
INKSOFT = RGBColor(0x5C, 0x4F, 0x40)
ACCENT = RGBColor(0x8C, 0x2F, 0x2B)
GOLD = RGBColor(0x9A, 0x7A, 0x33)
RULE = RGBColor(0xC9, 0xB9, 0x92)
SERIF = 'Palatino'
SANS = 'Helvetica Neue'

# pane number -> (title, caption) — keep in step with the gallery page
PANES = {
 '01': ('Overlay', 'The flagship reading pane: every word and phrase '
        'the dictionary knows, clickable — the longest known phrase '
        'lights up, the card answers below, the woodblock scan can '
        'follow along. Arrow keys walk the text; five display '
        'scripts including both pronunciation conventions.'),
 '02': ('Library', 'The preserved canon close at hand: the current '
        'Kangyur, Tengyur, and Sungbum releases, decoded catalog '
        'identities, English titles, one click into the Overlay — '
        'and a Maintenance menu that checks asianlegacylibrary.org '
        'for collection updates and installs them in place.'),
 '03': ('Scans', 'The woodblock tools for the loaded text: BDRC scan '
        'linking, whole-volume illustration galleries, and the '
        'four-layer folio view (scan · OCR · e-text · attested '
        'English).'),
 '04': ('Export', 'The publishing station: translation-prep format, '
        'the Pecha Maker (title folio, ya-yig, Degé measures, '
        'verse lineation, two-up imposition, batch mode), and '
        'print-ready Tibetan Unicode.'),
 '05': ('Draft', 'The working translator’s bench: source and draft '
        'side by side, the Evidence Ribbon following the cursor, '
        'established terms anchored, quotations detected, phrase '
        'memory, a live terminology guard as you type.'),
 '06': ('Manuscript', 'The composition surface: rich text with the '
        'corpus search beside it — the master’s attested English '
        'one “insert” away — and the publishing tools at hand.'),
 '07': ('Apparatus', 'Every published footnote and bibliography '
        'entry from the released volumes, sequential and '
        'searchable — the official tier only, reusable with one '
        'click.'),
 '08': ('Review', 'The overseer’s bench: register warnings, '
        'provisional-tier cautions, unmatched terms, collapsed '
        'distinctions. Guidance, never verdicts.'),
 '09': ('Align', 'The dictionary-building workflow reborn from the '
        'Hypercontext era: align Tibetan and English, harvest '
        'PENDING pairs — and collate two editions of a text into '
        'a variant apparatus.'),
 '10': ('Search', 'The classic Gofer grammar restored: OR and '
        'NEAR-within-N-lines across corpus, apparatus, Spotlight, '
        'and your own folders — folder results rolled up per file, '
        'one click opening the text at the hit.'),
 '11': ('Lookup', 'The stacked multi-dictionary: eight ways a query '
        'can land (headword, ACIP, English, pronunciation, ruled '
        'forms, community spellings, affix-stripped, browse), '
        'every layer labeled.'),
 '12': ('Sanskrit', 'The Sanskrit workbench: every notation, '
        'Whitney’s roots and grammar, Monier-Williams deep links, '
        'and the Mahāvyutpatti bridge.'),
 '13': ('Convert', 'Every writing system live as you type — ACIP ⇄ '
        'Wylie ⇄ Tibetan script, both pronunciation conventions '
        'side by side, and the full Tibetan calendar.'),
 '14': ('Analysis', 'The labeled-AI reading of one passage: engine '
        'pre-pass evidence, the model’s report, and a validation '
        'pass that re-checks every claim. Always banner-labeled.'),
 '15': ('Trainer', 'Learning to read, layer by layer: clause '
        'boundaries, particle roles, reading order, vocabulary, '
        'then the answer key — the master’s own English wherever '
        'the passage lives in his corpus.'),
 '16': ('Drills', 'Exercises that write themselves from the corpus — '
        'every answer is the master’s own text. Scrambles, '
        'fill-ins, particle choices, parallel reading, review.'),
 '17': ('Input', 'The input-center workstation, ACE reborn: type '
        'beside the scan, the image follows your cursor, '
        'predictive ACIP completion, OCR pre-fill, double-keying '
        'comparison.'),
 '18': ('OCR', 'Woodblock pages into text: line detection, '
        'recognition, illustration-candidate marking — every '
        'output headered OCR-DERIVED, review material only.'),
 '19': ('Propose', 'The team’s channel: offer a term, marking, '
        'pronunciation, or note — your name recorded, the passage '
        'attached as evidence.'),
 '20': ('Approval', 'The authority’s queue: every proposal ruled '
        'with provenance; register approvals apply in-app '
        'immediately; machine-seeded batches ruled in one '
        'considered act.'),
}


def add_run(p, text, size=None, color=None, bold=False, italic=False,
            font=None, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.name = font or SERIF
    if font:
        r._element.rPr.rFonts.set(qn('w:hAnsi'), font)
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    if caps:
        c = OxmlElement('w:caps')
        c.set(qn('w:val'), '1')
        r._element.rPr.append(c)
    if spacing:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(spacing))
        r._element.rPr.append(sp)
    return r


def para(doc, align=None, before=None, after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def pecha(doc):
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 14, 14)
    add_run(p, '─' * 20 + '  ', size=9, color=RULE)
    add_run(p, '❖', size=9, color=GOLD)
    add_run(p, '  ' + '─' * 20, size=9, color=RULE)


def main():
    doc = Document()
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), PAPER)
    doc.element.insert(0, bg)
    doc.settings.element.append(OxmlElement('w:displayBackgroundShape'))
    st = doc.styles['Normal']
    st.font.name = SERIF
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.28
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    img_width = Inches(6.7)

    # title page block
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 30, 8)
    add_run(p, 'Asian Legacy Library · ALL Translation Tool', size=8,
            color=GOLD, bold=True, font=SANS, caps=True, spacing=28)
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 0, 4)
    add_run(p, 'The Fifteen Panes', size=25, bold=True, color=ACCENT)
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    add_run(p, 'One application, one shared dictionary core — '
               'captured live, August 2026.', size=12, italic=True,
            color=INKSOFT)
    pecha(doc)

    shots = sorted(f for f in os.listdir(SHOTS) if f.endswith('.png'))
    for f in shots:
        num = f[:2]
        title, caption = PANES[num]
        p = para(doc, before=16, after=5)
        add_run(p, str(int(num)) + '  ', size=14, color=GOLD)
        add_run(p, title, size=14, bold=True, color=ACCENT)
        p.paragraph_format.keep_with_next = True
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(os.path.join(SHOTS, f),
                                  width=img_width)
        pic.paragraph_format.keep_with_next = True
        p = para(doc, before=4, after=10)
        add_run(p, caption, size=9.5, color=INKSOFT)

    pecha(doc)
    p = para(doc, WD_ALIGN_PARAGRAPH.CENTER, 6)
    add_run(p, 'Every screen above runs fully offline on a Mac, built '
               'on Geshe Michael Roach’s dictionary of 105,634 '
               'entries and his corpus of 42,199 aligned passages. The '
               'queue in pane 15 is real: 205 machine-derived '
               'pronunciations awaiting the authority’s ruling.',
            size=9.5, italic=True, color=INKSOFT)

    doc.save(OUT)
    print('wrote', OUT)


if __name__ == '__main__':
    main()
